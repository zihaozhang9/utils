#参考：https://www.cnblogs.com/wj-1314/p/9429816.html
#import pyocr
import importlib
import sys
import time
from docx import Document #pip install python-docx
 
importlib.reload(sys)
time1 = time.time()
# print("初始时间为：",time1)
 
import os.path
from pdfminer.pdfparser import  PDFParser,PDFDocument #py3:pip install pdfminer3k  ;pip install ply;pip install pyhcl
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal,LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
 
text_path = r'target2.pdf'
# text_path = r'photo-words.pdf'
 
def a_pdf2txt():
    '''解析PDF文本，并保存到TXT文件中'''
    fp = open(text_path,'rb')
    #用文件对象创建一个PDF文档分析器
    parser = PDFParser(fp)
    #创建一个PDF文档
    doc = PDFDocument()
    #连接分析器，与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
 
    #提供初始化密码，如果没有密码，就创建一个空的字符串
    doc.initialize()
 
    #检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        #创建PDF，资源管理器，来共享资源
        rsrcmgr = PDFResourceManager()
        #创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr,laparams=laparams)
        #创建一个PDF解释其对象
        interpreter = PDFPageInterpreter(rsrcmgr,device)
 
        #循环遍历列表，每次处理一个page内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            interpreter.process_page(page)
            #接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
            # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
            # 想要获取文本就获得对象的text属性，
            for x in layout:
                if(isinstance(x,LTTextBoxHorizontal)):
                    with open(r'results.txt','a',encoding='utf-8') as f:
                        results = x.get_text()
                        print(results)
                        if '标准答案' in results:continue
                        #import pdb;pdb.set_trace()
                        f.write(results  +"\n")  
 
def b_txt2word():
    document = Document()
    paragraph = document.add_paragraph('')

    str_lines = open(r'results.txt',encoding='utf-8').readlines()
    #import pdb;pdb.set_trace()
    for line in str_lines:
        if len(line.strip())<1:continue
        paragraph = document.add_paragraph('')
        if ('A' in line) or ('B' in line) or ('C' in line) or ('D' in line) :
            paragraph.add_run(line.strip())
        else:
            paragraph.add_run(line.strip()).bold = True
    document.save('results.docx')

if __name__ == '__main__':
    a_pdf2txt()
    b_txt2word()
    time2 = time.time()
    print("总共消耗时间为:",time2-time1)